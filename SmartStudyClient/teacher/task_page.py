import requests
from PyQt6.QtWidgets import QWidget, QFileDialog
from PyQt6.QtCore import pyqtSignal, QRect
from PyQt6.QtGui import QGuiApplication
from teacher.teacher_ui.task_page_form import Ui_task_page
from utils import utils
from docx import Document
from docx.shared import Cm
from PyQt6.QtGui import QPainter, QPixmap
from PyQt6.QtCore import Qt
from os import remove, listdir

class TaskPage(QWidget):
    task_page_finished = pyqtSignal(QRect)
    def __init__(self, token, subtopic_name, subtopic_id, topic_id, topic_name, geometry, session):
        super().__init__()
        utils.load_static(self)
        utils.load_user_settings(self)
        self.session=session
        self.token = token
        self.subtopic_name = subtopic_name
        self.subtopic_id = subtopic_id
        self.topic_id = topic_id
        self.topic_name = topic_name
        match self.user_settings['font-size']:
            case "small":
                self.window_size_index = 0
            case "medium":
                self.window_size_index = 1
            case "big":
                self.window_size_index = 2
        self.ui = Ui_task_page(self.static['user'][self.user_settings['font-size']])
        self.ui.setupUi(self)

        self.get_task()
        self.setGeometry(geometry)
        self.resize_window()

        self.ui.font_size_combo_box.currentIndexChanged.connect(self.resize_window)
        self.ui.back_to_topics_cmd_link_btn.clicked.connect(self.back_to_subtopic_page)
        if self.user_settings['maximized']:
            self.showMaximized()
        else:
            self.show()

    def resize_window(self):        
        selected_index = self.ui.font_size_combo_box.currentIndex()
        # When the user at the end do not want to change the font we need to correct the combo box which triggers the event again
        if selected_index == self.window_size_index:
            return
        screen_geometry = QGuiApplication.primaryScreen().geometry()
        screen_width = screen_geometry.width()
        screen_height = screen_geometry.height()
        match selected_index:
            case 0:
                self.user_settings['font-size']='small'
                self.window_size_index = 0
            case 1:
                if screen_width<self.static['user']['medium']['window-size']['width'] or screen_height<self.static['user']['medium']['window-size']['height']:
                    if not utils.too_small_screen():
                        self.ui.font_size_combo_box.setCurrentIndex(self.window_size_index)
                        return
                self.user_settings['font-size']='medium'
                self.window_size_index = 1
            case 2:
                if screen_width<self.static['user']['big']['window-size']['width'] or screen_height<self.static['user']['big']['window-size']['height']:
                    if not utils.too_small_screen():
                        self.ui.font_size_combo_box.setCurrentIndex(self.window_size_index)
                        return
                self.user_settings['font-size']='big'
                self.window_size_index = 2
        utils.save_settings(self)
        self.ui.static = self.static['user'][self.user_settings['font-size']]
        self.ui.screen_size(self)
        self.ui.task(self)

    def back_to_subtopic_page(self):
        self.user_settings['maximized'] = self.isMaximized()
        utils.save_settings(self)
        self.task_page_finished.emit(self.geometry())
        self.close()

    def get_task(self):
        url = f'{self.user_settings['SERVER_URL']}/smartstudy/'
        headers = {'Authorization': f'Token {self.token}'}
        json = {'action':'get_task', 'topic_id':self.topic_id, 'subtopic_id':self.subtopic_id}

        try:
            if url.startswith('https://127.0.0.1'):
                # Localhost SSL verification
                response = self.session.get(url, headers=headers, json=json)                
            else:
                response = self.session.get(url, headers=headers, json=json, verify=True)
            response_data = response.json()
            error = response_data.get('error')
            if error != None:
                utils.popup_window('Error', error)
                self.back_to_subtopic_page()
                return
            data = response_data.get('task')
            self.ui.data = data
            self.ui.task(self)
        except (requests.exceptions.RequestException, requests.exceptions.SSLError) as e:
            utils.popup_window('Error', f'Kérés elutasítva a szervertől: {self.user_settings['SERVER_URL']}/smartstudy/')
            self.back_to_subtopic_page()
            return

    def save_task(self):
        pixmap = QPixmap(self.ui.scene.sceneRect().size().toSize())
        pixmap.fill(Qt.GlobalColor.white)

        painter = QPainter(pixmap)
        self.ui.scene.render(painter)
        painter.end()

        pixmap.save("saved_tasks/scene_image.png")

        document = Document()

        document.add_picture('saved_tasks/scene_image.png', width=Cm(10))
        document.add_paragraph(self.ui.question_line_edit.toPlainText())
        document.add_paragraph(self.ui.answer_line_edit.text())        
        
        document.add_page_break()

        i = 1
        list_of_items = listdir('saved_tasks')
        while f"task_{i}.docx" in list_of_items:
            i+=1
        #document.save(f"saved_tasks/task_{i}.docx")
        remove("saved_tasks/scene_image.png")
        options = QFileDialog.Option(QFileDialog.Option.DontUseNativeDialog)
        file_name, _ = QFileDialog.getSaveFileName(self,f"feladat_{i}","","Word Documents (*.docx);;All Files (*)",options=options)
        if not file_name.endswith('.docx'):
            file_name += '.docx'
            document.save(file_name)
